import json
from pathlib import Path
from datetime import datetime
from docx import Document

research = (Path.cwd() / 'Research')
research.mkdir(exist_ok=True)

index = {"docs": []}

special_sections = [
    'Current Alignment',
    'Phase-by-Phase Gaps & Next Steps',
    'Cross-Cutting',
    'Immediate Next Steps (2–3 weeks)',
    'Milestones',
]

def get_headings(docx_path: Path):
    try:
        d = Document(str(docx_path))
    except Exception:
        return []
    heads = []
    for p in d.paragraphs:
        try:
            st = p.style.name if p.style else ''
        except Exception:
            st = ''
        if st and st.lower().startswith('heading') and p.text.strip():
            heads.append(p.text.strip())
    # unique in order
    seen = set(); uniq = []
    for h in heads:
        if h not in seen:
            seen.add(h); uniq.append(h)
    return uniq

for p in sorted(research.glob('*.docx')):
    stat = p.stat()
    heads = get_headings(p)
    has = { s: any(h.startswith(s) for h in heads) for s in special_sections }
    index["docs"].append({
        "name": p.name,
        "path": str(p.resolve()),
        "size": stat.st_size,
        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        "headings": heads[:20],
        "has_sections": has,
    })

(research / 'index.json').write_text(json.dumps(index, indent=2), encoding='utf-8')

# Build a small knowledge cache from redlined and summary docs
cache = {"sections": {}}
for fname in ['EAR_AI_Training_Proposal_redlined.docx', 'Research_Focus_Summary.docx']:
    path = research / fname
    if not path.exists():
        continue
    try:
        d = Document(str(path))
    except Exception:
        continue
    current = None; bucket = []
    def flush():
        if current:
            cache["sections"].setdefault(fname, {})[current] = [t for t in bucket if t]
    for para in d.paragraphs:
        txt = (para.text or '').strip()
        style = ''
        try:
            style = para.style.name or ''
        except Exception:
            pass
        if style.lower().startswith('heading') and txt:
            flush(); bucket = []; current = txt
        else:
            if txt:
                bucket.append(txt)
    flush()

(research / 'knowledge_cache.json').write_text(json.dumps(cache, indent=2), encoding='utf-8')
# Merge working_json_outline.json if present
try:
    wpath = research / 'working_json_outline.json'
    if wpath.exists():
        wpayload = json.loads(wpath.read_text(encoding='utf-8'))
        cache['working_json_outline'] = wpayload.get('working_json_outline', wpayload)
except Exception:
    pass
print('OK: index.json and knowledge_cache.json updated')


