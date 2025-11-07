from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

root = Path.cwd()
research = root / 'Research'
outline = research / 'Research_Manuscript_Outlines.docx'

if not outline.exists():
    raise SystemExit(f'Missing outline doc: {outline}')

# Files to link
files = [
    (research / 'Outlines_ISWC_SemWebJ.docx', 'ISWC / Semantic Web Journal Outlines', 'Ontology/SHACL, SPARQL lineage, KG quality metrics'),
    (research / 'Outlines_KDD_WWW_ACL_Industry.docx', 'KDD / WWW / ACL Industry Outlines', 'Hybrid RAG, SPARQL tool-use, SLOs/latency and ablations'),
    (research / 'Research_Path_Variants_By_Risk_Resourcing.docx', 'Path Variants by Risk/Resourcing', 'Low/Moderate/High tiers, DAL flag, timelines, prompts'),
    (research / 'Research_Focus_Summary.docx', 'Research Focus Summary', 'Themes + working JSON outline and variant paths'),
    (research / 'EAR_AI_Training_Proposal_redlined.docx', 'Proposal (Redlined, Updated)', 'Current Alignment, Gaps & Next Steps, Cross-Cutting, Immediate Steps, Milestones'),
    (research / 'Explainable Regulatory LLMs_ Current Landscape and Strategic Roadmap.docx', 'Strategic Roadmap', 'Background landscape and strategy statements'),
]


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF')
    rPr.append(u); rPr.append(color)
    new_run.append(rPr)

    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


doc = Document(str(outline))

# Add section header
h = doc.add_heading('Related Documents', level=2)

for path, title, note in files:
    if not path.exists():
        # Skip missing files silently
        continue
    p = doc.add_paragraph()
    file_url = path.resolve().as_uri()
    add_hyperlink(p, title, file_url)
    p.add_run(' — ' + note)

# Save
doc.save(str(outline))
print('Updated links in', outline)
