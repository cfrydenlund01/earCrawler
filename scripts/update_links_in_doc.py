import sys
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

if len(sys.argv) < 2:
    raise SystemExit('Usage: py scripts/update_links_in_doc.py <docx_path>')

docx_path = Path(sys.argv[1]).resolve()
research = docx_path.parent

files = [
    (research / 'Outlines_ISWC_SemWebJ.docx', 'ISWC / Semantic Web Journal Outlines', 'Ontology/SHACL, SPARQL lineage, KG quality metrics'),
    (research / 'Outlines_KDD_WWW_ACL_Industry.docx', 'KDD / WWW / ACL Industry Outlines', 'Hybrid RAG, SPARQL tool-use, SLOs/latency and ablations'),
    (research / 'Research_Path_Variants_By_Risk_Resourcing.docx', 'Path Variants by Risk/Resourcing', 'Low/Moderate/High tiers, DAL flag, timelines, prompts'),
    (research / 'Research_Focus_Summary.docx', 'Research Focus Summary', 'Themes + working JSON outline and variant paths'),
    (research / 'EAR_AI_Training_Proposal_redlined.docx', 'Proposal (Redlined, Updated)', 'Current Alignment, Gaps & Next Steps, Cross-Cutting, Immediate Steps, Milestones'),
    (research / 'Explainable Regulatory LLMs_ Current Landscape and Strategic Roadmap.docx', 'Strategic Roadmap', 'Background landscape and strategy statements'),
]


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF')
    rPr.append(u); rPr.append(color)
    new_run.append(rPr)

    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


doc = Document(str(docx_path))

doc.add_heading('Related Documents', level=2)
for path, title, note in files:
    if not path.exists():
        continue
    p = doc.add_paragraph()
    add_hyperlink(p, title, path.as_uri())
    p.add_run(' — ' + note)

doc.save(str(docx_path))
print('Updated links in', docx_path)
