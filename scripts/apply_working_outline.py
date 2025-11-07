import json
from pathlib import Path
from docx import Document

research = Path.cwd() / 'Research'
outline_file = research / 'working_json_outline.json'
summary_doc = research / 'Research_Focus_Summary.docx'
cache_file = research / 'knowledge_cache.json'

if not outline_file.exists():
    raise SystemExit(f'Missing {outline_file}')

payload = json.loads(outline_file.read_text(encoding='utf-8'))
pretty = json.dumps(payload, indent=2)

if summary_doc.exists():
    doc = Document(str(summary_doc))
    # Find the 'Working JSON Outline' heading and replace its following block
    start_idx = None
    for i, p in enumerate(doc.paragraphs):
        txt = (p.text or '').strip()
        sty = ''
        try:
            sty = p.style.name or ''
        except Exception:
            pass
        if txt == 'Working JSON Outline' and sty.lower().startswith('heading'):
            start_idx = i
            break
    if start_idx is not None:
        # remove until next heading or end
        j = start_idx + 1
        to_remove = []
        while j < len(doc.paragraphs):
            p = doc.paragraphs[j]
            sty = ''
            try:
                sty = p.style.name or ''
            except Exception:
                pass
            if sty.lower().startswith('heading'):
                break
            to_remove.append(p._element)
            j += 1
        for el in to_remove:
            parent = el.getparent();
            if parent is not None:
                parent.remove(el)
        # insert pretty JSON as a paragraph after heading
        doc.paragraphs[start_idx+1:start_idx+1]
        doc.paragraphs[start_idx]._element.addnext(doc._element.body._new_p())
        # Workaround: append at end then move the heading down
        doc.add_paragraph(pretty)
        doc.save(str(summary_doc))
    else:
        # no heading found; append a section
        doc.add_heading('Working JSON Outline', level=2)
        doc.add_paragraph(pretty)
        doc.save(str(summary_doc))

# Update knowledge cache with working_json_outline
try:
    cache = json.loads(cache_file.read_text(encoding='utf-8')) if cache_file.exists() else {}
except Exception:
    cache = {}
cache = dict(cache)
cache['working_json_outline'] = payload.get('working_json_outline', payload)
cache_file.write_text(json.dumps(cache, indent=2), encoding='utf-8')
# Also write to Research_Manuscript_Outlines.docx
manuscript_doc = research / 'Research_Manuscript_Outlines.docx'
if manuscript_doc.exists():
    try:
        mdoc = Document(str(manuscript_doc))
        mdoc.add_heading('Working JSON Outline', level=2)
        mdoc.add_paragraph(pretty)
        mdoc.save(str(manuscript_doc))
    except Exception:
        pass
print('Applied working_json_outline to summary and cache')

