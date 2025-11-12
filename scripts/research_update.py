import os, shutil, json
from pathlib import Path
from docx import Document

root = Path.cwd()
research_dir = root / "Research"
research_dir.mkdir(parents=True, exist_ok=True)

# Resolve source files (root or Research)
redlined_src = root / "EAR_AI_Training_Proposal_redlined.docx"
if not redlined_src.exists():
    redlined_src = research_dir / "EAR_AI_Training_Proposal_redlined.docx"
attached_src = (
    root / "Explainable Regulatory LLMs_ Current Landscape and Strategic Roadmap.docx"
)
if not attached_src.exists():
    attached_src = (
        research_dir
        / "Explainable Regulatory LLMs_ Current Landscape and Strategic Roadmap.docx"
    )

# If redlined doc missing entirely, create it fresh with new sections
if not redlined_src.exists():
    redlined_src = research_dir / "EAR_AI_Training_Proposal_redlined.docx"
    doc = Document()
    doc.add_heading("Current Alignment", level=1)
    for line in [
        "Phase A — Corpus Curation: loaders, transforms, analytics; deterministic fixtures present.",
        "Phase B — Knowledge Graph: ontology, SHACL, emit/load, Fuseki/Jena bootstrap, integrity gates, CI.",
        "Phase C — LLM Fine-Tuning: QLoRA scaffolding, sample datasets; config plumbing exists.",
        "Phase D — RAG Integration: retriever + agent modules; FastAPI facade + OpenAPI.",
        "Phase E — Evaluation: seeded harness and perf tests; needs task/metric expansion.",
        "Phase F — Explainability & Release: packaging, installer, SBOM, telemetry, RBAC, audit in place.",
    ]:
        doc.add_paragraph(line)
    doc.add_heading("Phase-by-Phase Gaps & Next Steps", level=1)
    for line in [
        "Phase A: finalize canonicalization/merge; provenance + redaction enforcement; snapshotting + jobs.",
        "Phase B: freeze v1 ontology/shapes; export profile verification; SPARQL perf budgets in CI.",
        "Phase C: dataset builders; experiment tracking/configs; model card + compliance checks; HPC templates.",
        "Phase D: unified retrieval + ranking; caching + tracing; agent tool-use for SPARQL.",
        "Phase E: define tasks/metrics and regression gates; dashboards; red-team and compliance tests.",
        "Phase F: explainability surfaces in API; signed artifacts; SDK + docs; Windows service ops.",
    ]:
        doc.add_paragraph(line)
    doc.add_heading("Cross-Cutting", level=1)
    for line in [
        "Security/RBAC/Audit: extend to RAG and model endpoints; deny-by-default.",
        "Observability: canaries + metrics for SPARQL latency, job lag, API availability.",
        "Privacy: enforce redaction across corpus, logs, telemetry, datasets.",
    ]:
        doc.add_paragraph(line)
    doc.add_heading("Immediate Next Steps (2–3 weeks)", level=1)
    for line in [
        "Lock Phase B baseline: freeze ontology/shapes; export profile hashes; SPARQL perf warmers.",
        "Phase A hardening: dedupe/merge + snapshots + operator run summaries.",
        "API + RAG touchpoint: align OpenAPI, add basic RAG endpoint stub with traces.",
        "Packaging & Ops: dry-run release scripts; signed artifacts + SBOM.",
    ]:
        doc.add_paragraph(line)
    doc.add_heading("Milestones", level=1)
    for line in [
        "M1: Windows-first end-to-end demo with CI gates.",
        "M2: Complete Phase B with export verification + API docs/SDK.",
        "M3: Phase C bootstrap with dataset builders and baselines.",
        "M4: Phase D/E with unified retrieval and continuous evaluation.",
        "M5: Phase F release with explainability and signed packages.",
    ]:
        doc.add_paragraph(line)
    doc.save(str(redlined_src))

# Ensure attached doc is in Research
if attached_src.exists() and attached_src.parent != research_dir:
    dst = research_dir / attached_src.name
    try:
        if dst.exists():
            dst.unlink()
        shutil.move(str(attached_src), str(dst))
        attached_src = dst
    except Exception:
        shutil.copy2(str(attached_src), str(dst))
        attached_src = dst

# If attached doc not present anywhere, abort
if not attached_src.exists():
    raise SystemExit(f"Missing attached DOCX: {attached_src}")

# Update redlined doc if it exists and is in Research
if redlined_src.exists():
    try:
        doc = Document(str(redlined_src))
        # remove prior addendum lines if any
        rm = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if "Addendum - Windows-First Interim Phase" in t or t.startswith(
                "Dependencies: Phases C-F"
            ):
                rm.append(p._element)
        for el in rm:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        doc.save(str(redlined_src))
    except Exception:
        pass

# Extract themes from attached document
from docx.opc.exceptions import PackageNotFoundError


def extract_themes(path: Path):
    try:
        d = Document(str(path))
    except PackageNotFoundError:
        return []
    headings = []
    for p in d.paragraphs:
        try:
            name = p.style.name if p.style else ""
        except Exception:
            name = ""
        if name and name.lower().startswith("heading") and p.text.strip():
            headings.append(p.text.strip())
    uniq = []
    seen = set()
    for h in headings:
        if h not in seen:
            seen.add(h)
            uniq.append(h)
    return uniq[:12]


themes = extract_themes(attached_src)
if not themes:
    try:
        d = Document(str(attached_src))
        themes = [p.text.strip() for p in d.paragraphs if p.text.strip()][:8]
    except Exception:
        themes = []

# JSON outlines
outline_core = {
    "objective": "Advance explainable regulatory LLMs with robust KG+RAG and reproducible evaluation for EAR/NSF corpora.",
    "research_axes": {
        "models": [
            {
                "type": "Llama 3.1 8B/70B",
                "finetune": ["LoRA", "QLoRA"],
                "explainability": ["saliency", "rationale extraction"],
            },
            {
                "type": "Mistral 7B",
                "finetune": ["QLoRA"],
                "explainability": ["attribution", "KG-path grounding"],
            },
            {"type": "Legal-specific BERT/LED", "use": ["reranking", "summarization"]},
        ],
        "kg_constructions": [
            {
                "scheme": "EAR ontology v1",
                "storage": "Jena TDB2",
                "reasoning": ["RDFS", "OWL Mini"],
            },
            {
                "scheme": "Entity-Paragraph graph",
                "links": ["cites", "wasDerivedFrom"],
                "profiles": ["export bundles"],
            },
        ],
        "retrieval": [
            {
                "index": "BM25+FAISS",
                "fusion": "reciprocal-rank fusion",
                "filters": ["entity type", "date"],
            },
            {
                "index": "Dense (E5/legal embeddings)",
                "hybrid": True,
                "cache": "feature+snippet",
            },
        ],
    },
    "workstreams": [
        {
            "name": "Corpus & Provenance",
            "tasks": ["curate JSONL", "dedupe/merge", "redaction", "snapshots"],
            "prompt": "Implement corpus build/validate/snapshot CLI with deterministic outputs and provenance checks.",
        },
        {
            "name": "KG & SPARQL",
            "tasks": ["freeze shapes", "export verification", "latency budgets"],
            "prompt": "Enforce export profiles and SPARQL perf warmers with CI gates.",
        },
        {
            "name": "RAG & Agent",
            "tasks": ["index builder", "ranking policy", "SPARQL tool-use", "tracing"],
            "prompt": "Expose a RAG endpoint aligned with OpenAPI; include sources, scores, lineage.",
        },
        {
            "name": "Eval & Explainability",
            "tasks": [
                "task suite",
                "benchmarks",
                "hallucination checks",
                "manuscript artifacts",
            ],
            "prompt": "Create evaluation gates and exportable figures/tables for publication.",
        },
    ],
    "manuscript_plan": {
        "tracks": [
            {
                "title": "Explainable RAG for EAR",
                "venue": "KDD/WWW/ACL Industry",
                "artifacts": ["datasets", "code", "appendix"],
            },
            {
                "title": "KG-anchored QA for Regulations",
                "venue": "ISWC/Semantic Web Journal",
            },
        ]
    },
    "acceptance": [
        "All tests pass",
        "Deterministic corpora",
        "Perf budgets met",
        "Reproducible eval",
    ],
    "milestones": [
        "M1 Windows-first demo",
        "M2 KG+API",
        "M3 Finetune baselines",
        "M4 RAG+Eval",
        "M5 Release",
    ],
}

variant_outlines = [
    {
        "title": "Hybrid RAG with KG Grounding",
        "hypotheses": [
            "Hybrid retrieval with entity filters improves precision",
            "Attribution to KG paths reduces hallucination",
        ],
        "methods": {
            "retrieval": "BM25 + dense (E5-legal) with RRF",
            "grounding": "SPARQL templates for lineage and citations",
            "generation": "Mistral 7B QLoRA with chain-of-thought (regulated)",
        },
        "datasets": ["EAR corpus", "NSF cases", "synthetic Q/A pairs"],
        "evaluation": {
            "metrics": [
                "hit@k",
                "nDCG@k",
                "Exact match",
                "BLEU/ROUGE",
                "Attribution score",
            ],
            "latency": {"p95_ms": 1500},
        },
        "risk_mitigations": ["privacy redaction", "denylist filters", "rate limits"],
        "gpt5_prompts": [
            "Implement indexer pipeline and hybrid retrieval with deterministic tests.",
            "Add SPARQL grounding step emitting source+path metadata in responses.",
        ],
    },
    {
        "title": "Reasoning over KG with Lightweight Generators",
        "hypotheses": ["OWL Mini inference enables fewer tokens for QA"],
        "methods": {
            "reasoning": "Fuseki OWL Mini endpoint for ASK/CONSTRUCT",
            "generator": "Small T5/LED summarizer for finalization",
        },
        "evaluation": {"metrics": ["accuracy", "token cost", "latency"]},
        "gpt5_prompts": [
            "Wire inference service and test ASK queries",
            "Compose answers from CONSTRUCT graphs",
        ],
    },
    {
        "title": "Instruction-Tuned Legal QA",
        "hypotheses": ["Domain instruction data reduces errors"],
        "methods": {
            "finetune": "QLoRA on Mistral/Llama",
            "safety": "policy-aware decoding",
        },
        "datasets": ["curated Q/A from EAR/NSF"],
        "gpt5_prompts": [
            "Dataset builder with redaction/licensing checks",
            "Trainer config + smoke eval",
        ],
    },
]

# Create Research_Focus_Summary.docx
summary_doc = research_dir / "Research_Focus_Summary.docx"
sdoc = Document()
sdoc.add_heading("Research Focus Summary", level=1)
if themes:
    sdoc.add_paragraph("Extracted Key Themes from attached document:")
    for t in themes:
        sdoc.add_paragraph(t)
sdoc.add_heading("Working JSON Outline", level=2)
sdoc.add_paragraph(json.dumps(outline_core, indent=2))
sdoc.add_heading("Variant Paths (JSON)", level=2)
sdoc.add_paragraph(json.dumps({"variants": variant_outlines}, indent=2))
sdoc.save(str(summary_doc))

# Create Research_Manuscript_Outlines.docx
outline_doc = research_dir / "Research_Manuscript_Outlines.docx"
odoc = Document()
odoc.add_heading("Manuscript-Oriented Research Outlines", level=1)
odoc.add_paragraph(
    "These outlines are structured for peer-review manuscript preparation."
)
plans = []
for v in variant_outlines:
    plan = {
        "title": v["title"],
        "concrete_steps": [
            "Define experimental design and fix seeds",
            "Build/validate indexes and corpora (deterministic)",
            "Implement endpoints/contracts and observability",
            "Run evaluation sweeps and export figures/tables",
            "Draft manuscript sections with methods/results",
        ],
        "instructions_for_gpt5": [
            "Follow acceptance criteria and do not break existing tests",
            "Add targeted unit tests for new logic",
            "Prefer Windows-first paths and offline fixtures",
            "Emit patch with minimal diff and clear structure",
        ],
        "acceptance": [
            "Existing + new tests pass",
            "Perf/latency budgets met",
            "Deterministic outputs",
            "Reproducible evaluation artifacts",
        ],
    }
    plans.append(plan)

odoc.add_heading("Variant Working Plans (JSON)", level=2)
odoc.add_paragraph(json.dumps({"plans": plans}, indent=2))
scaffold = {
    "base_header": [
        "You are a senior coding agent working in this repo.",
        "Keep changes minimal and deterministic; Windows-first.",
        "Validate with tests; avoid network in tests.",
    ],
    "work_package_template": {
        "goal": "...",
        "scope": ["..."],
        "key_files": ["path:line"],
        "deliverables": ["code", "tests", "docs"],
        "acceptance": ["..."],
        "validation": ["pytest -q", "py -m earCrawler.cli ..."],
        "non_goals": ["..."],
    },
}
odoc.add_heading("Prompt Scaffold (JSON)", level=2)
odoc.add_paragraph(json.dumps(scaffold, indent=2))
odoc.save(str(outline_doc))

print("OK: Research folder updated ->", research_dir)
