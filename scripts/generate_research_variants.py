import json
from pathlib import Path
from docx import Document

root = Path.cwd()
research_dir = root / "Research"
research_dir.mkdir(parents=True, exist_ok=True)

# 1) Path Variants by Risk/Resourcing
risk_variants = [
    {
        "tier": "Low-Risk / Compute-Light",
        "dal_required": False,
        "models": {
            "generator": "LED/T5-small summarizer",
            "reranker": "LegalBERT",
            "option": "Mistral 7B QLoRA (optional, single GPU)",
        },
        "kg": {
            "ontology": "EAR v1",
            "reasoning": ["none", "RDFS"],
            "storage": "Jena TDB2",
        },
        "retrieval": {
            "type": "BM25 + optional FAISS",
            "filters": ["entity type", "date"],
            "fusion": "RRF",
        },
        "explainability": ["KG lineage paths", "source snippets", "scores"],
        "timeline_weeks": 6,
        "tasks": [
            "Finalize corpus + snapshots",
            "Freeze shapes, export profiles + hashes",
            "Hybrid retrieval (BM25 base), add entity/date filters",
            "RAG endpoint with lineage + source scores",
            "Evaluation: hit@k, latency P95, attribution score",
            "Draft manuscript (methods, system, results)",
        ],
        "acceptance": [
            "Deterministic corpora",
            "Latency P95 < 1500ms",
            "Attribution ≥ baseline",
        ],
        "gpt5_prompts": [
            "Implement BM25 retriever with entity/date filters (earCrawler/rag/retriever.py) and tests.",
            "Add lineage emission in API responses (service/api_server/routers/entities.py).",
        ],
    },
    {
        "tier": "Moderate-Risk / Single-GPU",
        "dal_required": False,
        "models": {
            "generator": "Mistral 7B QLoRA",
            "embeddings": "E5-legal-base",
            "reranker": "LegalBERT",
        },
        "kg": {
            "ontology": "EAR v1",
            "reasoning": ["RDFS", "OWL Mini"],
            "storage": "Jena TDB2",
        },
        "retrieval": {
            "type": "Hybrid BM25 + dense",
            "fusion": "RRF",
            "cache": "feature+snippet",
        },
        "explainability": ["KG path grounding", "reranker rationales"],
        "timeline_weeks": 8,
        "tasks": [
            "Dense index build + hybrid fusion",
            "QLoRA adapter fine-tune on curated Q/A",
            "SPARQL tool-use for grounding",
            "Evaluation sweeps: retrieval + QA accuracy",
            "Manuscript ablations: dense vs hybrid vs BM25",
        ],
        "acceptance": [
            "QA EM improves ≥ X%",
            "Retrieval nDCG@10 ≥ baseline",
            "P95 < 2000ms",
        ],
        "gpt5_prompts": [
            "Add dense indexing and fusion policy with tests (earCrawler/rag/retriever.py).",
            "Wire QLoRA trainer configs and smoke eval (earCrawler/agent/mistral_agent.py).",
        ],
    },
    {
        "tier": "High-Risk / DAL-HPC",
        "dal_required": True,
        "models": {
            "generator": "Llama 3.1 70B QLoRA/LoRA",
            "embeddings": "E5-large-legal",
        },
        "kg": {
            "ontology": "EAR v1",
            "reasoning": ["OWL Mini"],
            "storage": "TDB2 inference service",
        },
        "retrieval": {
            "type": "Hybrid + learned reranking",
            "fusion": "weighted",
            "cache": "embedding+snippet",
        },
        "explainability": ["counterfactual grounding", "path saliency"],
        "timeline_weeks": 12,
        "tasks": [
            "HPC job templates + data staging",
            "Multi-epoch QLoRA with evaluation checkpoints",
            "Large-scale indexing + canary/perf dashboards",
            "Manuscript large-scale results + error analysis",
        ],
        "acceptance": [
            "Significant QA gains vs moderate tier",
            "Stable latency under load",
            "Reproducible runs on DAL",
        ],
        "gpt5_prompts": [
            "Generate SLURM/TACC job templates and dataset staging scripts.",
            "Add large-scale evaluation harness with exportable tables.",
        ],
    },
]

risk_doc = research_dir / "Research_Path_Variants_By_Risk_Resourcing.docx"
risk = Document()
risk.add_heading("Path Variants by Risk/Resourcing", level=1)
risk.add_paragraph(json.dumps({"variants": risk_variants}, indent=2))
risk.save(str(risk_doc))

# 2) Venue-specific outlines: ISWC/SemWebJ
iswc = {
    "target": ["ISWC", "Semantic Web Journal"],
    "thesis": "KG-grounded QA for regulations with explainable SPARQL lineage and ontology-driven validation.",
    "contributions": [
        "Versioned EAR ontology + SHACL shapes",
        "KG ↔ SPARQL templates ↔ QA alignment with lineage",
        "Reproducible pipelines and export profiles",
    ],
    "methods": {
        "kg": {
            "ontology": "EAR v1",
            "reasoning": ["RDFS", "OWL Mini"],
            "validation": "SHACL",
        },
        "grounding": ["template-based SPARQL", "path extraction"],
        "explainability": ["lineage graphs", "attribution"],
        "evaluation": {
            "kg_quality": ["SHACL conforms", "isomorphism"],
            "qa": ["accuracy", "attribution"],
            "perf": ["SPARQL latency"],
        },
    },
    "artifacts": ["TTL, shapes, bundles, SPARQL templates", "Open-source code"],
    "gpt5_prompts": [
        "Freeze shapes and add migration notes (earCrawler/kg/shapes.ttl).",
        "Implement ASK/CONSTRUCT tests with OWL Mini (tests/kg/*).",
    ],
}

iswc_doc = research_dir / "Outlines_ISWC_SemWebJ.docx"
d1 = Document()
d1.add_heading("ISWC / Semantic Web Journal Outlines", level=1)
d1.add_paragraph(json.dumps(iswc, indent=2))
d1.save(str(iswc_doc))

# 3) Venue-specific outlines: KDD/WWW/ACL Industry
kdd = {
    "target": ["KDD", "WWW", "ACL Industry"],
    "thesis": "Production-grade explainable RAG over regulatory corpora with strict SLOs and deterministic pipelines.",
    "contributions": [
        "Windows-first deterministic ingestion + KG",
        "Hybrid retrieval with SPARQL tool-use",
        "Observability stack with latency/canary SLOs",
    ],
    "methods": {
        "retrieval": {"hybrid": True, "fusion": "RRF", "reranking": "LegalBERT"},
        "generation": {
            "model": "Mistral/Llama QLoRA",
            "guardrails": ["policy-aware decoding"],
        },
        "ops": {"latency_budgets_ms": {"api_p95": 800, "sparql_p95": 1500}},
    },
    "evaluation": {
        "metrics": ["hit@k", "nDCG@k", "EM/F1", "attribution"],
        "perf": ["P95 latency", "throughput"],
        "ablation": ["bm25 vs dense vs hybrid", "with/without tool-use"],
    },
    "gpt5_prompts": [
        "Implement hybrid indexer + caching (earCrawler/rag/retriever.py).",
        "Add API traces and latency gauges (service/api_server/logging_integration.py).",
    ],
}

kdd_doc = research_dir / "Outlines_KDD_WWW_ACL_Industry.docx"
d2 = Document()
d2.add_heading("KDD / WWW / ACL Industry Outlines", level=1)
d2.add_paragraph(json.dumps(kdd, indent=2))
d2.save(str(kdd_doc))

print("Generated:", risk_doc.name, iswc_doc.name, kdd_doc.name)
